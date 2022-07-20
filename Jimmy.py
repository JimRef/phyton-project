from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# Profile Picture
document.add_picture('dp.jpg', width=Inches(0.5))

# Your details Name cp number email
name = input("What is your name? ")
speak('Hi' + name + 'How are you today')
speak('What is your phone number?')
phone_number = input("What is your phone number? ")
email = input("What is your email? ")

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# About Me
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself '))

# WorkExperience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter The Name of Your Previous Company ')
from_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + "\n").italic = True

experience_details = input('Describe your experience at ' + company + ' ')

p.add_run(experience_details)

# More experience
while True:
    more_experience = input('Do you have more experience? Yes or No ')
    if more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter The Name of Your Previous Company ')
        from_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + "\n").italic = True

        experience_details = input('Describe your experience at ' + company + ' ')

        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')
skills = input('Enter your Skill ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    more_skills = input('Do you have more skills? Yes or No ')
    if more_skills.lower() == 'yes':

        skills = input('Enter your Skill ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

# Reference
document.add_heading('Reference')
p = document.add_paragraph()

colleague_name = input('Enter a Name of your Colleague/Friends ')
colleague_contact_number = input('Enter the Contact Number of your Colleague ')


p.add_run(colleague_name + "\n").bold = True
p.add_run(colleague_contact_number).italic = True

# More Reference
while True:
    more_colleague = input('Do you have more Colleague/Friends? Yes or No ')
    if more_colleague == 'yes':
        p = document.add_paragraph()

        colleague_name = input('Enter a Name of your Colleague/Friends ')
        colleague_contact_number = input('Enter the Contact Number of your Colleague ')

        p.add_run(colleague_name + "\n").bold = True
        p.add_run(colleague_contact_number).italic = True
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Resume Generated Using Jimmy Refugio Codes'


document.save('Resume.docx')
